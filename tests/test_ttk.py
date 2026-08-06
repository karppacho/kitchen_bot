"""Тесты генерации ТТК: сборка контекста и рендер .docx. Офлайн, без сети и LLM.

Форма документа задаётся образцом `TTK_template.pdf` и собирается скриптом
`build_ttk_template.py`. Шаблон .docx — генерируемый артефакт (в .gitignore),
поэтому тест рендера собирает его во временный каталог сам и не зависит от
того, лежит ли готовый шаблон в корне проекта.
"""
import runpy
from decimal import Decimal
from pathlib import Path

import pytest
from docx import Document

from src.ttk import builder
from src.ttk.builder import _ttk_number_from_dish, build_ttk_context, render_ttk
from tests.test_calc import make_data


# Поля, которых в новой форме нет. Тест-сторож: если кто-то вернёт их в контекст,
# шаблон о них не знает и они молча потеряются.
REMOVED_KEYS = {"org_name", "director_position", "approval_date", "tr_ts_number"}


@pytest.fixture(scope="module")
def template_path(tmp_path_factory) -> Path:
    """Собирает шаблон скриптом во временный каталог."""
    out_dir = tmp_path_factory.mktemp("ttk_template")
    script = Path(__file__).resolve().parents[1] / "build_ttk_template.py"
    import os
    cwd = os.getcwd()
    try:
        os.chdir(out_dir)
        runpy.run_path(str(script), run_name="__main__")
    finally:
        os.chdir(cwd)
    path = out_dir / "TTK_template.docx"
    assert path.exists(), "build_ttk_template.py не создал шаблон"
    return path


def _docx_text(path: Path) -> str:
    """Весь текст документа: абзацы + ячейки таблиц."""
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


# ---------- контекст ----------

def test_context_has_no_legal_fields():
    """Реквизиты ООО, дата и ТР ТС из формы убраны — их не должно быть и в контексте."""
    context, _ = build_ttk_context(make_data(), "T001")
    assert REMOVED_KEYS.isdisjoint(context), (
        f"в контексте остались поля, которых нет в шаблоне: "
        f"{REMOVED_KEYS & set(context)}"
    )


def test_context_recipe_and_output():
    """Рецептура — только основные ингредиенты, упаковка в ТТК не попадает."""
    context, meta = build_ttk_context(make_data(), "T001")
    names = [i["name"] for i in context["ingredients"]]
    assert names == ["Тортилья", "Салат"]
    assert "Коробка" not in names            # упаковка не входит в рецептуру
    assert context["dish_output_g"] == "160"  # 60 + 100, брутто в выход не идёт
    assert context["dish_name"] == "Тест-ролл"
    assert meta["has_composition"] is True


def test_context_brutto_accounts_for_losses():
    """Брутто считается с потерями, нетто — как в ТТК."""
    d = make_data()
    d.ingredients[3].losses_cutting = Decimal("0.2")   # салат: нарезка 20%
    context, _ = build_ttk_context(d, "T001")
    salad = next(i for i in context["ingredients"] if i["name"] == "Салат")
    assert salad["netto"] == "100"
    assert float(salad["brutto"]) == pytest.approx(125.0, abs=0.01)


def test_document_uses_iiko_names():
    """В карту идёт «Короткое для айки» — по ней работают технологи в iiko."""
    d = make_data()
    d.ingredients[1].pos_name = "Тортилья ПФ"
    d.ingredients[3].pos_name = "Салат айсберг пф"
    context, meta = build_ttk_context(d, "T001")

    assert [i["name"] for i in context["ingredients"]] == ["Тортилья ПФ", "Салат айсберг пф"]
    # А в подсказку для LLM идут обычные имена: про «салат» писать проще, чем про «пф»
    assert [i["name"] for i in meta["ingredients"]] == ["Тортилья", "Салат"]
    assert not any("Короткое для айки" in w for w in meta["warnings"])


def test_missing_iiko_name_falls_back_and_warns():
    """Колонка E пустая → берём обычное имя, но обязательно предупреждаем.

    Пустая ячейка в официальной карте хуже, чем неточное имя, а молча подменять
    название нельзя: технолог должен знать, что имя не сверено с iiko.
    """
    d = make_data()
    d.ingredients[1].pos_name = "Тортилья ПФ"
    d.ingredients[3].pos_name = ""            # у салата не заполнено
    context, meta = build_ttk_context(d, "T001")

    assert [i["name"] for i in context["ingredients"]] == ["Тортилья ПФ", "Салат"]
    warn = [w for w in meta["warnings"] if "Короткое для айки" in w]
    assert len(warn) == 1
    assert "Салат" in warn[0]
    assert "Тортилья" not in warn[0]          # у неё имя заполнено, она ни при чём


def test_iiko_name_whitespace_only_is_treated_as_empty():
    d = make_data()
    d.ingredients[1].pos_name = "   "
    context, _ = build_ttk_context(d, "T001")
    assert context["ingredients"][0]["name"] == "Тортилья"


def test_context_text_fields_start_empty():
    """Техпроцесс и органолептику заполняет LLM-слой перед рендером."""
    context, _ = build_ttk_context(make_data(), "T001")
    for key in ("tech_process", "organoleptic_appearance", "organoleptic_color",
                "organoleptic_taste_smell", "organoleptic_consistency"):
        assert context[key] == ""


def test_context_unknown_dish():
    assert build_ttk_context(make_data(), "НЕТ-ТАКОГО") is None


def test_ttk_number_from_dish():
    assert _ttk_number_from_dish("B001") == "1"
    assert _ttk_number_from_dish("B132") == "132"
    assert _ttk_number_from_dish("без-цифр") == "без-цифр"


# ---------- рендер ----------

def test_render_produces_docx(template_path, tmp_path, monkeypatch):
    """Полный проход: контекст → рендер → в файле нужные разделы и данные."""
    monkeypatch.setattr(builder, "TEMPLATE_PATH", template_path)
    context, _ = build_ttk_context(make_data(), "T001")
    context["tech_process"] = "Прогреть в пресс-гриле 30 секунд."
    context["organoleptic_appearance"] = "Ровный рулет, разрезан наискось."
    context["organoleptic_color"] = "Золотистая тортилья, зелёный салат."
    context["organoleptic_taste_smell"] = "Свежий вкус, запах гриля."
    context["organoleptic_consistency"] = "Мягкая лепёшка, хрустящий салат."

    out = render_ttk(context, tmp_path / "ttk.docx")
    assert out.exists()

    text = _docx_text(out)
    # Четыре раздела новой формы — и ничего из старой
    assert "1. РЕЦЕПТУРА" in text
    assert "2. ТЕХНОЛОГИЧЕСКИЙ ПРОЦЕСС" in text
    assert "3. ПОКАЗАТЕЛИ КАЧЕСТВА И БЕЗОПАСНОСТИ" in text
    assert "4. ПИЩЕВАЯ И ЭНЕРГЕТИЧЕСКАЯ ЦЕННОСТЬ" in text
    for gone in ("УТВЕРЖДАЮ", "ОБЛАСТЬ ПРИМЕНЕНИЯ", "ТРЕБОВАНИЯ К СЫРЬЮ",
                 "Разработал", "ТР ТС"):
        assert gone not in text, f"в документе осталось «{gone}» из старой формы"

    # Данные подставились
    assert "Тест-ролл" in text
    assert "Тортилья" in text
    assert "Выход блюда:" in text
    assert "160" in text
    assert "Прогреть в пресс-гриле" in text
    assert "Ровный рулет" in text
    # Плейсхолдеры не должны просочиться в готовый документ
    assert "{{" not in text and "{%" not in text


# ---------- пакетная генерация ----------

def _batch_data():
    """База с блюдами разных статусов; у одного нет состава."""
    from src.data.models import Dish
    d = make_data()
    base = d.ttk_by_dish["T001"]
    d.dishes = {
        "B1": Dish(id="B1", name="Новинка А", category="Пицца", status="разработка",
                   price_menu=Decimal("300")),
        "B2": Dish(id="B2", name="Новинка Б", category="Бургер", status="разработка",
                   price_menu=Decimal("300")),
        "B3": Dish(id="B3", name="Без состава", category="Пицца", status="разработка",
                   price_menu=Decimal("300")),
        "B4": Dish(id="B4", name="Старое", category="Пицца", status="активное",
                   price_menu=Decimal("300")),
    }
    d.ttk_by_dish = {
        k: [r.model_copy(update={"dish_id": k}) for r in base]
        for k in ("B1", "B2", "B4")
    }
    return d


def test_batch_preview_lists_only_requested_status(monkeypatch):
    import src.llm.client as client
    d = _batch_data()
    monkeypatch.setattr(client, "get_data", lambda: d)

    r = client._tool_generate_ttk_batch({"status": "новинки"})
    assert r["count"] == 2                       # B3 без состава не считается
    assert "Новинка А" in r["display"]
    assert "Старое" not in r["display"]          # активное не попало
    assert "Без состава" in r["display"]         # но о пропуске сказали
    assert "file_paths" not in r                 # на превью файлы не делаем


def test_batch_refuses_above_limit(monkeypatch):
    """Пачка больше лимита — отказ: на каждую карту два запроса к модели."""
    import src.llm.client as client
    from src.data.models import Dish
    d = make_data()
    base = d.ttk_by_dish["T001"]
    d.dishes = {
        f"B{i}": Dish(id=f"B{i}", name=f"Блюдо {i}", category="Пицца",
                      status="разработка", price_menu=Decimal("300"))
        for i in range(client.MAX_TTK_BATCH + 3)
    }
    d.ttk_by_dish = {
        k: [r.model_copy(update={"dish_id": k}) for r in base] for k in d.dishes
    }
    monkeypatch.setattr(client, "get_data", lambda: d)

    r = client._tool_generate_ttk_batch({"status": "разработка", "confirm": True})
    assert "error" in r
    assert str(client.MAX_TTK_BATCH) in r["error"]


def test_batch_unknown_status(monkeypatch):
    import src.llm.client as client
    monkeypatch.setattr(client, "get_data", lambda: _batch_data())
    assert "error" in client._tool_generate_ttk_batch({"status": "абракадабра"})


def test_batch_one_failure_does_not_kill_the_rest(monkeypatch, template_path):
    """Сбой одной карты не должен ронять пачку — остальные всё равно приходят."""
    import src.llm.client as client
    d = _batch_data()
    monkeypatch.setattr(client, "get_data", lambda: d)
    monkeypatch.setattr(builder, "TEMPLATE_PATH", template_path)

    calls = {"n": 0}

    def _fake_render(context, meta, tech_hint=""):
        calls["n"] += 1
        if calls["n"] == 1:
            raise RuntimeError("модель не ответила")
        return template_path

    monkeypatch.setattr(client, "_fill_and_render_ttk", _fake_render)

    r = client._tool_generate_ttk_batch({"status": "разработка", "confirm": True})
    assert r["made"] == 1
    assert r["failed"] == 1
    assert len(r["file_paths"]) == 1
    assert "модель не ответила" in r["display"]


def test_render_without_template_raises(tmp_path, monkeypatch):
    """Нет шаблона — понятная ошибка с подсказкой, а не молчаливый сбой."""
    monkeypatch.setattr(builder, "TEMPLATE_PATH", tmp_path / "нет-файла.docx")
    context, _ = build_ttk_context(make_data(), "T001")
    with pytest.raises(FileNotFoundError, match="build_ttk_template"):
        render_ttk(context, tmp_path / "out.docx")


def test_render_ingredient_rows_match_composition(template_path, tmp_path, monkeypatch):
    """Строк рецептуры ровно столько, сколько ингредиентов — не фиксированные 8."""
    monkeypatch.setattr(builder, "TEMPLATE_PATH", template_path)
    context, _ = build_ttk_context(make_data(), "T001")
    out = render_ttk(context, tmp_path / "ttk2.docx")

    doc = Document(str(out))
    recipe = doc.tables[0]
    # 2 строки шапки + по строке на ингредиент + строка «Выход блюда»
    assert len(recipe.rows) == 2 + len(context["ingredients"]) + 1
