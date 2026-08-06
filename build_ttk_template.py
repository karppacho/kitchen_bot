"""
Создаёт Word-шаблон ТТК для сети "Тим Кук" в формате docxtpl.

Запуск:  python build_ttk_template.py  →  TTK_template.docx в корне проекта.
Файл в .gitignore: это генерируемый артефакт, эталон формы — TTK_template.pdf.

Версия 3 (август 2026): форма переведена на новый образец шефа (TTK_template.pdf).
Осталось четыре раздела вместо восьми:
  1. Рецептура   2. Технологический процесс
  3. Показатели качества и безопасности   4. Пищевая и энергетическая ценность

Убраны (их нет в новой форме): шапка «УТВЕРЖДАЮ» с реквизитами ООО и датой,
«Область применения», «Требования к сырью», «Требования к оформлению, реализации
и хранению», «Предусмотренное применение», ссылки на ТР ТС, строка «Разработал».

Версия 2: заметные границы, серая шапка, отступы, keep_with_next на заголовках,
неразрывные таблицы КБЖУ, повтор шапки при переносе.

Опечатки исходных файлов исправлены ("БЕЗОПАСТНОСТИ", "ФРАНЦУСКОГО", "заетм") —
в новой форме опечатка в заголовке раздела 3 тоже есть, но мы её не переносим.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# В образце шефа (TTK_template.pdf) шрифт выглядит как Arial, но это дефолт
# браузера при печати, а не выбор сети. Для официального документа оставляем
# Times New Roman — меняется здесь одной строкой.
BASE_FONT = 'Times New Roman'


# ---------- ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ----------

def set_cell_borders(cell, sz='8'):
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn('w:tcBorders')):
        tc_pr.remove(old)
    tc_borders = OxmlElement('w:tcBorders')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_padding(cell, top=40, bottom=40, left=120, right=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for direction, value in (('top', top), ('bottom', bottom),
                              ('left', left), ('right', right)):
        node = OxmlElement(f'w:{direction}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement('w:cantSplit')
    tr_pr.append(cant_split)


def set_paragraph_keep_with_next(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement('w:keepNext')
    pPr.append(keep)


def style_cell(cell, text, *, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT,
               valign=WD_ALIGN_VERTICAL.CENTER, shading=None, font=BASE_FONT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    cell.vertical_alignment = valign
    set_cell_borders(cell)
    set_cell_padding(cell)
    if shading:
        set_cell_shading(cell, shading)


def add_paragraph(doc, text, *, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=0, keep_with_next=False,
                  font=BASE_FONT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if keep_with_next:
        set_paragraph_keep_with_next(p)
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    return p


def make_table_header_repeat(row):
    """Делает строку таблицы повторяющейся шапкой при переносе таблицы."""
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement('w:tblHeader'))


def merge_and_style(table, r1, c1, r2, c2, text, **style_kwargs):
    """Объединить диапазон ячеек и оформить получившуюся.

    В новой форме объединение нужно трижды: шапка «расход сырья…» на две колонки
    и заголовки обеих таблиц КБЖУ. style_cell вызывается ПОСЛЕ merge — иначе
    заливка и границы теряются при слиянии.
    """
    cell = table.cell(r1, c1).merge(table.cell(r2, c2))
    style_cell(cell, text, **style_kwargs)
    return cell


# ---------- КОНСТАНТЫ ----------

HEADER_FILL = 'D9D9D9'
SECTION_GAP_BEFORE = 12
SECTION_GAP_AFTER = 6
TABLE_GAP_BEFORE = 4
TABLE_GAP_AFTER = 8


def make_kbju_table(doc, title, p_var, f_var, c_var, k_var, title_extra=None):
    """Таблица КБЖУ: строка-титул + шапка колонок + значения.

    title       — «пищевая и энергетическая ценность на 100 грамм блюда»;
    title_extra — если задано, титул занимает 3 колонки, а в четвёртой стоит это
                  значение (в новой форме там выход блюда для таблицы «на 1 блюдо»).

    Все строки cant_split — таблица не рвётся между страницами.
    """
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_widths = [Cm(3.5), Cm(3.5), Cm(4), Cm(5.5)]
    for col_idx, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[col_idx].width = width

    title_row = table.rows[0]
    set_row_cant_split(title_row)
    make_table_header_repeat(title_row)
    if title_extra is None:
        merge_and_style(table, 0, 0, 0, 3, title, bold=True,
                        align=WD_ALIGN_PARAGRAPH.CENTER, shading=HEADER_FILL)
    else:
        merge_and_style(table, 0, 0, 0, 2, title, bold=True,
                        align=WD_ALIGN_PARAGRAPH.CENTER, shading=HEADER_FILL)
        style_cell(table.cell(0, 3), title_extra, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, shading=HEADER_FILL)

    hdr = table.rows[1]
    set_row_cant_split(hdr)
    make_table_header_repeat(hdr)
    for cell, text in zip(hdr.cells,
                          ['белки, г', 'жиры, г', 'углеводы, г',
                           'энергетическая ценность, кКал']):
        style_cell(cell, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                   shading=HEADER_FILL)

    val_row = table.rows[2]
    set_row_cant_split(val_row)
    for cell, text in zip(val_row.cells, [p_var, f_var, c_var, k_var]):
        style_cell(cell, text, align=WD_ALIGN_PARAGRAPH.CENTER)


# ---------- ОСНОВНАЯ ФУНКЦИЯ ----------

def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    style = doc.styles['Normal']
    style.font.name = BASE_FONT
    style.font.size = Pt(11)

    # === ЗАГОЛОВОК ===
    add_paragraph(doc, 'ТЕХНИКО-ТЕХНОЛОГИЧЕСКАЯ КАРТА № {{ ttk_number }}',
                  bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER,
                  keep_with_next=True)
    add_paragraph(doc, '{{ dish_name }}', bold=True, size=13,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # === 1. РЕЦЕПТУРА ===
    add_paragraph(doc, '1. РЕЦЕПТУРА', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=SECTION_GAP_BEFORE, space_after=SECTION_GAP_AFTER,
                  keep_with_next=True)

    # Шапка в два яруса: «расход сырья…» объединён на брутто+нетто,
    # а «наименование сырья и продуктов» — на обе строки шапки.
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_widths = [Cm(9), Cm(3.75), Cm(3.75)]
    for col_idx, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[col_idx].width = width

    for row in table.rows[:2]:
        set_row_cant_split(row)
        make_table_header_repeat(row)

    merge_and_style(table, 0, 0, 1, 0, 'наименование сырья и продуктов',
                    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, shading=HEADER_FILL)
    merge_and_style(table, 0, 1, 0, 2, 'расход сырья и продуктов на 1 порцию, г',
                    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, shading=HEADER_FILL)
    style_cell(table.cell(1, 1), 'брутто', bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, shading=HEADER_FILL)
    style_cell(table.cell(1, 2), 'нетто', bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, shading=HEADER_FILL)

    # Строк ровно столько, сколько ингредиентов у блюда: docxtpl-цикл по строкам.
    # Фиксированные 8 строк бумажной формы не годятся — состав бывает и длиннее.
    for_row = table.add_row()
    set_row_cant_split(for_row)
    style_cell(for_row.cells[0], '{%tr for ing in ingredients %}')
    style_cell(for_row.cells[1], '')
    style_cell(for_row.cells[2], '')

    ing_row = table.add_row()
    set_row_cant_split(ing_row)
    style_cell(ing_row.cells[0], '{{ ing.name }}')
    style_cell(ing_row.cells[1], '{{ ing.brutto }}', align=WD_ALIGN_PARAGRAPH.CENTER)
    style_cell(ing_row.cells[2], '{{ ing.netto }}', align=WD_ALIGN_PARAGRAPH.CENTER)

    endfor_row = table.add_row()
    set_row_cant_split(endfor_row)
    style_cell(endfor_row.cells[0], '{%tr endfor %}')
    style_cell(endfor_row.cells[1], '')
    style_cell(endfor_row.cells[2], '')

    output_row = table.add_row()
    set_row_cant_split(output_row)
    style_cell(output_row.cells[0], 'Выход блюда:', bold=True,
               align=WD_ALIGN_PARAGRAPH.RIGHT, shading=HEADER_FILL)
    style_cell(output_row.cells[1], '—', bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, shading=HEADER_FILL)
    style_cell(output_row.cells[2], '{{ dish_output_g }}', bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, shading=HEADER_FILL)
    for col_idx, width in enumerate(col_widths):
        output_row.cells[col_idx].width = width

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(TABLE_GAP_AFTER)

    # === 2. ТЕХНОЛОГИЧЕСКИЙ ПРОЦЕСС ===
    add_paragraph(doc, '2. ТЕХНОЛОГИЧЕСКИЙ ПРОЦЕСС', bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=SECTION_GAP_BEFORE, space_after=SECTION_GAP_AFTER,
                  keep_with_next=True)
    add_paragraph(
        doc,
        'Подготовка сырья производится в соответствии с рекомендациями Сборника '
        'технологических нормативов для предприятий общественного питания и '
        'технологическими рекомендациями для импортного сырья.',
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4,
    )
    add_paragraph(
        doc,
        'Продукт готовить под конкретный заказ. {{ tech_process }}',
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    # === 3. ПОКАЗАТЕЛИ КАЧЕСТВА И БЕЗОПАСНОСТИ ===
    # В форме шефа заголовок с опечаткой («БЕЗОПАСТНОСТИ») — не переносим.
    add_paragraph(doc, '3. ПОКАЗАТЕЛИ КАЧЕСТВА И БЕЗОПАСНОСТИ', bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=SECTION_GAP_BEFORE, space_after=SECTION_GAP_AFTER,
                  keep_with_next=True)

    org_table = doc.add_table(rows=5, cols=2)
    org_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    org_col_widths = [Cm(5), Cm(11.5)]
    for col_idx, width in enumerate(org_col_widths):
        for row in org_table.rows:
            row.cells[col_idx].width = width

    org_rows_data = [
        ('Наименование показателей', 'Характеристика показателей', True),
        ('Внешний вид', '{{ organoleptic_appearance }}', False),
        ('Цвет', '{{ organoleptic_color }}', False),
        ('Вкус и запах', '{{ organoleptic_taste_smell }}', False),
        ('Консистенция', '{{ organoleptic_consistency }}', False),
    ]
    for row_idx, (label, value, is_header) in enumerate(org_rows_data):
        row = org_table.rows[row_idx]
        set_row_cant_split(row)
        if is_header:
            style_cell(row.cells[0], label, bold=True,
                       align=WD_ALIGN_PARAGRAPH.CENTER, shading=HEADER_FILL)
            style_cell(row.cells[1], value, bold=True,
                       align=WD_ALIGN_PARAGRAPH.CENTER, shading=HEADER_FILL)
            make_table_header_repeat(row)
        else:
            style_cell(row.cells[0], label)
            style_cell(row.cells[1], value, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(TABLE_GAP_AFTER)

    # === 4. ПИЩЕВАЯ И ЭНЕРГЕТИЧЕСКАЯ ЦЕННОСТЬ ===
    add_paragraph(doc, '4. ПИЩЕВАЯ И ЭНЕРГЕТИЧЕСКАЯ ЦЕННОСТЬ', bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=SECTION_GAP_BEFORE, space_after=SECTION_GAP_AFTER,
                  keep_with_next=True)

    make_kbju_table(
        doc, 'пищевая и энергетическая ценность на 100 грамм блюда',
        '{{ kbju_per_100g.белки }}', '{{ kbju_per_100g.жиры }}',
        '{{ kbju_per_100g.углеводы }}', '{{ kbju_per_100g.ккал }}',
    )

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(TABLE_GAP_AFTER)

    # Во второй таблице в шапке справа стоит выход блюда — как в форме шефа.
    make_kbju_table(
        doc, 'пищевая и энергетическая ценность на 1 блюдо',
        '{{ kbju_per_portion.белки }}', '{{ kbju_per_portion.жиры }}',
        '{{ kbju_per_portion.углеводы }}', '{{ kbju_per_portion.ккал }}',
        title_extra='{{ dish_output_g }}',
    )

    doc.save('TTK_template.docx')
    print('Сохранён: TTK_template.docx')


if __name__ == '__main__':
    main()
